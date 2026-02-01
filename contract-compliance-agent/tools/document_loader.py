"""
Document Loader for Contract Compliance Agent.

Handles loading and text extraction from various document formats
including PDF, DOCX, and plain text files.
"""

from pathlib import Path
from typing import Optional

from models.schemas import DocumentMetadata


# =============================================================================
# Document Loader
# =============================================================================

class DocumentLoader:
    """
    Loads and extracts text from contract documents.
    
    Supports multiple file formats with graceful fallback
    when optional dependencies are not available.
    """
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}
    
    def __init__(self):
        """Initialize the document loader and check available parsers."""
        self._pdf_available = self._check_pdf_support()
        self._docx_available = self._check_docx_support()
    
    @staticmethod
    def _check_pdf_support() -> bool:
        """Check if PDF parsing is available."""
        try:
            import pypdf
            return True
        except ImportError:
            return False
    
    @staticmethod
    def _check_docx_support() -> bool:
        """Check if DOCX parsing is available."""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    # -------------------------------------------------------------------------
    # Main Loading Methods
    # -------------------------------------------------------------------------
    
    def load(self, file_path: str | Path) -> tuple[str, DocumentMetadata]:
        """
        Load a document and extract its text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, document_metadata)
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type is not supported
        """
        path = Path(file_path)
        
        # Validate file exists
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")
        
        # Validate file type
        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {suffix}. "
                f"Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # Extract text based on file type
        if suffix == ".pdf":
            text = self._load_pdf(path)
        elif suffix in {".docx", ".doc"}:
            text = self._load_docx(path)
        else:
            text = self._load_text(path)
        
        # Build metadata
        metadata = DocumentMetadata(
            filename=path.name,
            file_type=suffix,
            file_size=path.stat().st_size,
            word_count=len(text.split()),
            character_count=len(text),
            language="en"  # Default; could be enhanced with detection
        )
        
        return text, metadata
    
    def load_text(self, file_path: str | Path) -> str:
        """
        Load a document and return only the text content.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text content
        """
        text, _ = self.load(file_path)
        return text
    
    # -------------------------------------------------------------------------
    # Format-Specific Loaders
    # -------------------------------------------------------------------------
    
    def _load_pdf(self, path: Path) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        if not self._pdf_available:
            raise ImportError(
                "PDF support requires pypdf. "
                "Install with: pip install pypdf"
            )
        
        import pypdf
        
        text_parts = []
        page_count = 0
        
        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            page_count = len(reader.pages)
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    def _load_docx(self, path: Path) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        if not self._docx_available:
            raise ImportError(
                "DOCX support requires python-docx. "
                "Install with: pip install python-docx"
            )
        
        import docx
        
        doc = docx.Document(path)
        
        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n\n".join(paragraphs)
    
    def _load_text(self, path: Path) -> str:
        """
        Load a plain text file.
        
        Args:
            path: Path to text file
            
        Returns:
            File content
        """
        # Try common encodings
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                with open(path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # Fallback: read as bytes and decode with replacement
        with open(path, "rb") as f:
            return f.read().decode("utf-8", errors="replace")
    
    # -------------------------------------------------------------------------
    # Utility Methods
    # -------------------------------------------------------------------------
    
    def get_supported_formats(self) -> dict[str, bool]:
        """
        Get information about supported formats and their availability.
        
        Returns:
            Dict mapping format to availability status
        """
        return {
            ".txt": True,
            ".md": True,
            ".pdf": self._pdf_available,
            ".docx": self._docx_available,
            ".doc": self._docx_available,
        }
    
    @staticmethod
    def is_supported(file_path: str | Path) -> bool:
        """
        Check if a file type is supported.
        
        Args:
            file_path: Path to check
            
        Returns:
            True if file type is supported
        """
        path = Path(file_path)
        return path.suffix.lower() in DocumentLoader.SUPPORTED_EXTENSIONS


# =============================================================================
# Convenience Functions
# =============================================================================

def load_document(file_path: str | Path) -> tuple[str, DocumentMetadata]:
    """
    Load a document and extract its text.
    
    Convenience function that creates a loader and processes
    a single document.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Tuple of (text, metadata)
    """
    loader = DocumentLoader()
    return loader.load(file_path)


def load_text(file_path: str | Path) -> str:
    """
    Load a document and return just the text.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Extracted text content
    """
    loader = DocumentLoader()
    return loader.load_text(file_path)


def extract_contract_sections(text: str) -> dict[str, str]:
    """
    Attempt to identify and extract contract sections.
    
    Uses common section headers to segment the contract.
    
    Args:
        text: Full contract text
        
    Returns:
        Dict mapping section names to their content
    """
    # Common contract section patterns
    section_patterns = [
        r"(?:^|\n)(\d+\.?\s*[A-Z][A-Za-z\s]+)(?:\n|:)",  # "1. DEFINITIONS"
        r"(?:^|\n)(ARTICLE\s+[IVX\d]+[:\s]+[A-Za-z\s]+)",  # "ARTICLE I: SCOPE"
        r"(?:^|\n)([A-Z][A-Z\s]{3,})(?:\n|:)",  # "CONFIDENTIALITY"
    ]
    
    import re
    
    sections = {}
    current_section = "PREAMBLE"
    current_content = []
    
    for line in text.split("\n"):
        # Check if line is a section header
        is_header = False
        for pattern in section_patterns:
            if match := re.match(pattern, line):
                # Save previous section
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                
                # Start new section
                current_section = match.group(1).strip()
                current_content = []
                is_header = True
                break
        
        if not is_header:
            current_content.append(line)
    
    # Save final section
    if current_content:
        sections[current_section] = "\n".join(current_content).strip()
    
    return sections
